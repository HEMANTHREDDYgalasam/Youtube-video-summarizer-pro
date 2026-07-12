# ==============================================================================
# utils.py — Helper utilities for YouTube Video Summarizer Pro
# Author  : Hemanth Reddy
# Purpose : Video ID extraction, transcript fetching, thumbnail URL building,
#           AI summary generation, and file export helpers.
# ==============================================================================

from __future__ import annotations

import io
import re
import time
from typing import Optional

import requests
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import (
    TranscriptsDisabled,
    NoTranscriptFound,
    VideoUnavailable,
    VideoUnplayable,
    AgeRestricted,
    RequestBlocked,
    IpBlocked,
    CouldNotRetrieveTranscript,
)


# ─────────────────────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────────────────────

# ── Gemini model ──────────────────────────────────────────────────────────────
# Single source-of-truth for the model name used throughout the app.
# "gemini-2.0-flash" is the GA stable release — fast, cost-efficient, and
# available to all Google AI Studio users (free & paid).
# To upgrade: change this one line. Do NOT hardcode model names elsewhere.
GEMINI_MODEL: str = "gemini-2.0-flash"

# ── Summary prompts ───────────────────────────────────────────────────────────
SUMMARY_PROMPTS: dict[str, str] = {
    "Short (~100 words)": (
        "Summarize the following YouTube video transcript in approximately 100 words. "
        "Be concise, capture the core topic, and highlight the single most important insight. "
        "Write in clear, engaging prose. Do NOT use bullet points."
    ),
    "Medium (~250 words)": (
        "Summarize the following YouTube video transcript in approximately 250 words. "
        "Cover the main points, key arguments, and important takeaways. "
        "Use a structured paragraph format with a brief intro, body, and conclusion. "
        "Do NOT use bullet points."
    ),
    "Detailed (~500 words)": (
        "Provide a comprehensive summary of the following YouTube video transcript in approximately 500 words. "
        "Include: a strong introduction, detailed coverage of all major topics discussed, "
        "key insights and examples mentioned, and a meaningful conclusion. "
        "Use well-structured paragraphs with smooth transitions. Do NOT use bullet points."
    ),
}


# ─────────────────────────────────────────────────────────────────────────────
# Video ID Extraction
# ─────────────────────────────────────────────────────────────────────────────

def extract_video_id(url: str) -> Optional[str]:
    """
    Extract the YouTube video ID from various URL formats.

    Supported formats:
      - https://www.youtube.com/watch?v=VIDEO_ID
      - https://youtu.be/VIDEO_ID
      - https://youtube.com/shorts/VIDEO_ID
      - https://www.youtube.com/embed/VIDEO_ID
      - https://www.youtube.com/v/VIDEO_ID

    Args:
        url: A YouTube URL string.

    Returns:
        The 11-character video ID string, or None if not found.
    """
    patterns = [
        r"(?:v=|/v/|/embed/|youtu\.be/|/shorts/)([A-Za-z0-9_-]{11})",
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Thumbnail
# ─────────────────────────────────────────────────────────────────────────────

def get_thumbnail_url(video_id: str) -> str:
    """
    Return the highest-resolution available thumbnail URL for a YouTube video.

    Tries maxresdefault first, falls back to hqdefault.

    Args:
        video_id: The 11-character YouTube video ID.

    Returns:
        A publicly accessible thumbnail URL string.
    """
    maxres = f"https://img.youtube.com/vi/{video_id}/maxresdefault.jpg"
    try:
        response = requests.head(maxres, timeout=5)
        # YouTube returns a 120×90 placeholder for missing maxres thumbnails
        if response.status_code == 200 and int(response.headers.get("Content-Length", 0)) > 5000:
            return maxres
    except requests.RequestException:
        pass

    return f"https://img.youtube.com/vi/{video_id}/hqdefault.jpg"


# ─────────────────────────────────────────────────────────────────────────────
# Transcript Fetching
# ─────────────────────────────────────────────────────────────────────────────

def fetch_transcript(video_id: str) -> tuple[str, str]:
    """
    Fetch and concatenate the transcript for a YouTube video.

    Compatible with youtube-transcript-api >= 1.0 (instance-based API).
    Attempts English first (manual, then auto-generated), then falls back
    to any available language in the transcript list.

    Supports:
      - Normal videos with English transcripts
      - Multilingual videos (falls back to first available language)
      - Auto-generated subtitles (YouTube's ASR captions)
      - Manually uploaded subtitles

    Args:
        video_id: The 11-character YouTube video ID.

    Returns:
        A tuple of (transcript_text, language_code).

    Raises:
        ValueError: With a user-friendly message describing the failure reason.
    """
    try:
        # ── v1.x API: must instantiate the class, then call .list() ──
        ytt_api = YouTubeTranscriptApi()
        transcript_list = ytt_api.list(video_id)

        # Priority 1: Manually created English transcript
        try:
            transcript = transcript_list.find_manually_created_transcript(["en"])
        except Exception:
            # Priority 2: Auto-generated English transcript
            try:
                transcript = transcript_list.find_generated_transcript(["en"])
            except Exception:
                # Priority 3: First available transcript in any language
                try:
                    transcript = next(iter(transcript_list))
                except StopIteration:
                    raise ValueError(
                        "No transcript is available for this video. "
                        "The creator may not have enabled subtitles."
                    )

        # Fetch the actual transcript data
        fetched = transcript.fetch()

        # FetchedTranscript is iterable; each item is a FetchedTranscriptSnippet
        # dataclass with a .text attribute (youtube-transcript-api >= 1.0)
        text_parts: list[str] = []
        for snippet in fetched:
            # FetchedTranscriptSnippet dataclass: .text, .start, .duration
            if hasattr(snippet, "text"):
                text_parts.append(snippet.text)
            elif isinstance(snippet, dict):
                # Defensive fallback for any future API changes
                text_parts.append(snippet.get("text", ""))

        full_text = " ".join(text_parts).strip()

        if not full_text:
            raise ValueError("The transcript appears to be empty.")

        return full_text, transcript.language_code

    except VideoUnavailable:
        raise ValueError(
            "⚠️ This video is unavailable. It may be private or deleted."
        )
    except VideoUnplayable:
        raise ValueError(
            "⚠️ This video is unplayable (may be region-restricted or require login)."
        )
    except AgeRestricted:
        raise ValueError(
            "⚠️ This video is age-restricted. Transcripts cannot be fetched without authentication."
        )
    except TranscriptsDisabled:
        raise ValueError(
            "⚠️ Transcripts are disabled for this video by the creator."
        )
    except NoTranscriptFound:
        raise ValueError(
            "⚠️ No transcript could be found for this video. "
            "Try a different video with subtitles/closed captions enabled."
        )
    except RequestBlocked:
        raise ValueError(
            "⚠️ YouTube blocked the transcript request. "
            "This may be a temporary rate-limit. Please try again in a moment."
        )
    except IpBlocked:
        raise ValueError(
            "⚠️ Your IP address has been blocked by YouTube. "
            "Try using a VPN or a different network."
        )
    except CouldNotRetrieveTranscript as exc:
        raise ValueError(f"⚠️ Could not retrieve transcript: {exc}") from exc
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"⚠️ Failed to fetch transcript: {exc}") from exc


# ─────────────────────────────────────────────────────────────────────────────
# AI Summary Generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_summary(
    transcript: str,
    mode: str,
    api_key: str,
) -> tuple[str, float]:
    """
    Generate an AI-powered summary of a video transcript using Gemini.

    Uses the google-genai SDK (google.genai) — the latest supported API.

    Args:
        transcript: The full transcript text.
        mode: One of the keys in SUMMARY_PROMPTS.
        api_key: A valid Google Gemini API key.

    Returns:
        A tuple of (summary_text, elapsed_seconds).

    Raises:
        ValueError: With a user-friendly message on API or generation failure.
    """
    try:
        from google import genai  # type: ignore

        client = genai.Client(api_key=api_key)

        system_instruction = SUMMARY_PROMPTS.get(mode, SUMMARY_PROMPTS["Medium (~250 words)"])

        # Truncate extremely long transcripts to avoid token limits (~30 000 chars ≈ ~7 500 tokens)
        max_transcript_chars = 30_000
        if len(transcript) > max_transcript_chars:
            transcript = transcript[:max_transcript_chars] + "\n\n[Transcript truncated for length]"

        prompt = f"{system_instruction}\n\n---\nTRANSCRIPT:\n{transcript}\n---"

        start_time = time.perf_counter()

        # Try multiple fallback models to bypass quota limits or regional/account model restrictions
        models_to_try = [GEMINI_MODEL, "gemini-flash-latest", "gemini-2.0-flash-lite"]
        unique_models = []
        for m in models_to_try:
            if m not in unique_models:
                unique_models.append(m)

        last_exc = None
        for model_name in unique_models:
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt,
                )
                summary = response.text.strip()
                if not summary:
                    raise ValueError("Gemini returned an empty response. Please try again.")
                
                elapsed = time.perf_counter() - start_time
                return summary, elapsed
            except Exception as exc:
                last_exc = exc
                error_msg = str(exc)
                
                # If API Key is fundamentally invalid, fail fast as trying another model won't help
                if "API_KEY_INVALID" in error_msg or "api key" in error_msg.lower():
                    raise ValueError(
                        "❌ Invalid Google API Key. Please check your key and try again."
                    ) from exc
                
                # Log to streamlit console if running interactively
                print(f"[Model Fallback] Model '{model_name}' failed: {error_msg[:100]}...")
                continue

        # If we reached here, all models failed
        if last_exc:
            error_msg = str(last_exc)
            if "404" in error_msg or "NOT_FOUND" in error_msg or "no longer available" in error_msg.lower():
                raise ValueError(
                    f"❌ None of the tested Gemini models could be found or accessed on your account. "
                    "Please check your Google AI Studio plan or billing details."
                ) from last_exc
            if "quota" in error_msg.lower() or "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                raise ValueError(
                    "❌ API quota exceeded on all candidate models. Please wait a moment and try again, "
                    "or check your Google AI Studio limits at https://aistudio.google.com/"
                ) from last_exc
            raise ValueError(f"❌ AI generation failed: {last_exc}") from last_exc
        else:
            raise ValueError("❌ AI generation failed: No models available to try.")

    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"❌ AI generation failed: {exc}") from exc


# ─────────────────────────────────────────────────────────────────────────────
# Text Statistics
# ─────────────────────────────────────────────────────────────────────────────

def compute_stats(text: str) -> dict[str, int | float]:
    """
    Compute word count and estimated reading time for a text block.

    Args:
        text: Any string of text.

    Returns:
        A dict with keys:
          - word_count (int)
          - char_count (int)
          - reading_time_seconds (float)  — assumes 200 WPM average reading speed
    """
    words = text.split()
    word_count = len(words)
    char_count = len(text)
    reading_time_seconds = (word_count / 200) * 60  # 200 WPM

    return {
        "word_count": word_count,
        "char_count": char_count,
        "reading_time_seconds": reading_time_seconds,
    }


def format_reading_time(seconds: float) -> str:
    """
    Format reading time in seconds into a human-readable string.

    Args:
        seconds: Reading time in seconds.

    Returns:
        A human-readable string like "< 1 min" or "2 min 30 sec".
    """
    if seconds < 60:
        return f"< 1 min"
    minutes = int(seconds // 60)
    remaining_secs = int(seconds % 60)
    if remaining_secs == 0:
        return f"{minutes} min"
    return f"{minutes} min {remaining_secs} sec"


# ─────────────────────────────────────────────────────────────────────────────
# Export Helpers
# ─────────────────────────────────────────────────────────────────────────────

def export_as_txt(summary: str, video_url: str, mode: str) -> bytes:
    """
    Encode a summary as UTF-8 plain text bytes.

    Args:
        summary: The summary text.
        video_url: The original YouTube URL.
        mode: The summary mode label.

    Returns:
        UTF-8 encoded bytes.
    """
    header = (
        "=" * 60 + "\n"
        "YouTube Video Summarizer Pro\n"
        "Created by Hemanth Reddy\n"
        "=" * 60 + "\n\n"
        f"Video URL : {video_url}\n"
        f"Mode      : {mode}\n"
        f"Generated : {time.strftime('%Y-%m-%d %H:%M:%S')}\n"
        + "-" * 60 + "\n\n"
    )
    content = header + summary + "\n"
    return content.encode("utf-8")


def export_as_pdf(summary: str, video_url: str, mode: str) -> bytes:
    """
    Generate a professionally formatted PDF from a summary.

    Args:
        summary: The summary text.
        video_url: The original YouTube URL.
        mode: The summary mode label.

    Returns:
        Raw PDF bytes.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    styles = getSampleStyleSheet()

    # FIX: Use HexColor objects instead of raw hex strings for reportlab
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=20,
        textColor=HexColor("#6C63FF"),
        spaceAfter=6,
        alignment=TA_CENTER,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=HexColor("#888888"),
        spaceAfter=4,
        alignment=TA_CENTER,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontSize=9,
        textColor=HexColor("#555555"),
        spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=11,
        leading=18,
        textColor=HexColor("#222222"),
        alignment=TA_JUSTIFY,
        spaceAfter=12,
    )

    story = [
        Paragraph("YouTube Video Summarizer Pro", title_style),
        Paragraph("Created by Hemanth Reddy", subtitle_style),
        Spacer(1, 0.2 * inch),
        Paragraph(f"<b>Video URL:</b> {video_url}", meta_style),
        Paragraph(f"<b>Summary Mode:</b> {mode}", meta_style),
        Paragraph(f"<b>Generated:</b> {time.strftime('%Y-%m-%d %H:%M:%S')}", meta_style),
        Spacer(1, 0.3 * inch),
    ]

    # Split paragraphs to preserve structure
    for para in summary.split("\n\n"):
        clean = para.strip().replace("\n", " ")
        if clean:
            story.append(Paragraph(clean, body_style))

    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph("─" * 60, meta_style))
    story.append(Paragraph("Generated with YouTube Video Summarizer Pro", subtitle_style))

    doc.build(story)
    return buffer.getvalue()


def export_as_docx(summary: str, video_url: str, mode: str) -> bytes:
    """
    Generate a professionally formatted DOCX from a summary.

    Args:
        summary: The summary text.
        video_url: The original YouTube URL.
        mode: The summary mode label.

    Returns:
        Raw DOCX bytes.
    """
    document = Document()

    # ── Styles ──────────────────────────────────────────────
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title_para = document.add_heading("YouTube Video Summarizer Pro", level=0)
    title_para.alignment = 1  # CENTER
    title_run = title_para.runs[0]
    title_run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)

    # Subtitle
    sub = document.add_paragraph("Created by Hemanth Reddy")
    sub.alignment = 1
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sub.runs[0].font.size = Pt(10)

    document.add_paragraph()

    # Metadata table
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    meta = [
        ("Video URL", video_url),
        ("Summary Mode", mode),
        ("Generated", time.strftime("%Y-%m-%d %H:%M:%S")),
    ]
    for i, (label, value) in enumerate(meta):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    document.add_paragraph()

    # Section heading
    section_heading = document.add_heading("Summary", level=1)
    section_heading.runs[0].font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)

    # Summary body
    for para_text in summary.split("\n\n"):
        clean = para_text.strip()
        if clean:
            p = document.add_paragraph(clean)
            p.alignment = 3  # JUSTIFY
            p.paragraph_format.space_after = Pt(8)

    document.add_paragraph()

    # Footer
    footer_para = document.add_paragraph("─" * 50)
    credit = document.add_paragraph("Generated with YouTube Video Summarizer Pro")
    credit.alignment = 1
    credit.runs[0].font.size = Pt(9)
    credit.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
